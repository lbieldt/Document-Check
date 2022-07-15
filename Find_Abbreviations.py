import re
from typing import Counter
import docx2txt
from docx import Document

import tkinter as tk
from tkinter import filedialog


root = tk.Tk()
root.withdraw()

file_path = filedialog.askopenfilename() #Load file path with GUI popup
text = docx2txt.process(file_path)
# open connection to Word Document, read as txt
#text = docx2txt.process("C:\Python\Acronyms\Test.docx") #Static path to test



# Test abbreviation expression here: https://regex101.com/r/A1SJoe/16
rx = r"\b[A-Z](?=([&.]?))(?:\1[A-Z])+\b" # Looks for all the double Cap letters or special chars
s = text # Load text string
abbreviations =  [x.group() for x in re.finditer(rx, s)] #Save abbreviations in a list
abbr_sorted = sorted(abbreviations) #Sort list alfabetically

count_abbr = Counter(abbr_sorted) #Count all the unique items => Counter({'IPS': 14, 'HSS': 2, 'RCCS': 2, 'RB': 1})
unique_abbr = list(dict.fromkeys(abbr_sorted)) #List all the unique items => ['HSS', 'HSS', 'IPS', 'IPS', 'IPS', 'IPS', 'IPS', 'IPS', 'IPS', 'IPS', 'IPS', 
print(count_abbr["IPS"])
# Create a new word document
document = Document()
# Create a table with all the unique Abbreviation
table = document.add_table(rows=1, cols=2) # Create table with two columns
hdr_cells = table.rows[0].cells

hdr_cells[0].text = 'Abbreviation'
hdr_cells[1].text = 'Occurrence'
for abbr in unique_abbr:
    row_cells = table.add_row().cells
    #occurance = abbr
    row_cells[0].text = abbr
    row_cells[1].text = str(count_abbr[abbr])
    #print(count_abbr[abbr]) add_run('bold').bold = True

table.style ="Table Grid"

file_path_save = filedialog.asksaveasfilename() #Load file path with GUI popup
document.save(file_path_save)

print(count_abbr)
print(unique_abbr)
